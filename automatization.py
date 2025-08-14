import pandas as pd
from docx import Document

ruta = "ORIGINAL/StockEquiposTecnologicos-2025.xlsx"

df = pd.read_excel(ruta, header=2).dropna(axis=1, how="all").dropna(how="all")

numero = 5
fila = df[df["N°"] == numero]

datos = dict()
if not fila.empty:
    datos = fila.iloc[0].to_dict()
else:
    print("No encontrado")

colaborador= "Alejandro"
dni = "1"


ruta_word = "ORIGINAL/ActaEntregaLaptopCropsP.docx"
doc= Document(ruta_word)

colaborador_id = f"{{{{COLABORADOR}}}}"   
dni_id = f"{{{{DNI}}}}" 

for i in doc.paragraphs:
    
    if colaborador_id in i.text and dni_id in i.text:
        i.text = i.text.replace(colaborador_id, colaborador)
        i.text = i.text.replace(dni_id,dni)


if 'N°' in datos:
    datos.pop('N°') 


datos['CODIGO'] = datos['ALIAS']
datos['S/N']=datos['NÚMERO DE SERIE']
datos['MODELO DEL CARGADOR'] = datos['NÚMERO DE SERIE CARGADOR']

for tabla in doc.tables:
    for fila_tabla in tabla.rows:
        for celda in fila_tabla.cells:
            for clave, valor in datos.items():
                marcador = f"{{{{{clave}}}}}"

               
                if marcador in celda.text :
                
                    celda.text = celda.text.replace(marcador, str(valor))


salida = f"COPIA/Acta_Laptop_1.docx"
doc.save(salida)
print(f"Documento generado: {salida}")